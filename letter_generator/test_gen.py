from docx import Document

# Création d'un nouveau document Word
doc = Document()

# Ajouter un titre ou un objet (facultatif)
doc.add_heading("Objet : Demande d'information", level=1)

# Ajouter le contenu du template avec des balises dynamiques
doc.add_paragraph("Bonjour {nom},")
doc.add_paragraph("")
doc.add_paragraph("Nous vous remercions de l'intérêt que vous portez à nos services.")
doc.add_paragraph("")
doc.add_paragraph("Nous souhaitons vous informer que votre adresse enregistrée est :")
doc.add_paragraph("{adresse}")
doc.add_paragraph("")
doc.add_paragraph("Pour toute question, n'hésitez pas à nous contacter à l'adresse suivante :")
doc.add_paragraph("{email}")
doc.add_paragraph("")
doc.add_paragraph("Cordialement,")
doc.add_paragraph("L'équipe de Mon Application")

# Sauvegarder le document dans un fichier
output_filename = "modele_lettre.docx"
doc.save(output_filename)
print(f"Le template a été généré et sauvegardé sous le nom '{output_filename}'.")
