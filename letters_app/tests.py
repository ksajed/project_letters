from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Créer un nouveau document
doc = Document()

# --- Personnalisation de l'en-tête ---
section = doc.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.text = "Nom de l'Entreprise"
header_paragraph.style = doc.styles['Header']
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Optionnel : Ajouter un logo dans l'en-tête (décommentez et fournissez le chemin du fichier)
# header.add_picture('chemin/vers/logo.png', width=Inches(1.0))

# --- Ajout de la date ---
p_date = doc.add_paragraph()
p_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run_date = p_date.add_run("Date : {date}")
run_date.font.size = Pt(10)

# --- Ajout de l'objet de la lettre ---
p_objet = doc.add_paragraph()
p_objet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_objet = p_objet.add_run("Objet : Demande d'information")
run_objet.font.size = Pt(14)
run_objet.bold = True

doc.add_paragraph()  # Ligne vide

# --- Corps de la lettre ---
# Salutations initiales
p_salutation = doc.add_paragraph("Madame, Monsieur,")
p_salutation.space_after = Pt(12)

# Paragraphe d'introduction
p_intro = doc.add_paragraph()
p_intro.add_run("Nous vous remercions de l'intérêt que vous portez à nos services. "
                "Nous vous contactons afin de vous fournir toutes les informations nécessaires "
                "concernant votre demande.").font.size = Pt(12)

doc.add_paragraph()  # Ligne vide

# Détails du contact (balises dynamiques)
p_details = doc.add_paragraph()
p_details.add_run("Nom du Client : {nom}\n").font.size = Pt(12)
p_details.add_run("Adresse : {adresse}\n").font.size = Pt(12)
p_details.add_run("Email : {email}\n").font.size = Pt(12)

doc.add_paragraph()  # Ligne vide

# Paragraphe de conclusion
p_conclusion = doc.add_paragraph()
p_conclusion.add_run("Nous restons à votre disposition pour toute information complémentaire. "
                      "Veuillez agréer, Madame, Monsieur, l'expression de nos salutations distinguées.").font.size = Pt(12)

doc.add_paragraph()  # Ligne vide

# Signature
p_signature = doc.add_paragraph("L'équipe de [Nom de l'Entreprise]")
p_signature.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
p_signature.runs[0].font.size = Pt(12)
p_signature.runs[0].italic = True

# --- Personnalisation du pied de page ---
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = "Adresse de l'Entreprise | Téléphone: 01 23 45 67 89 | Email: contact@entreprise.com"
footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer_paragraph.runs[0].font.size = Pt(10)

# --- Enregistrer le document ---
output_filename = "modele_lettre_professionnel.docx"
doc.save(output_filename)
print(f"Le template professionnel a été généré et sauvegardé sous le nom '{output_filename}'.")
